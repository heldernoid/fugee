"""agent/tools/doc_generator.py — Phase 5 document package (PDF + editable Word).

Generates the document package from real session data. The **personal statement
is drafted by the LLM** (see agent/drafting.py) — an in-depth, first-person
narrative with ``[placeholders]`` the person completes — and every document is
exported as both an **editable Word (.docx)** file and a print-ready **PDF**.
Country facts come from the curated data via country_lookup; nothing is
fabricated (CLAUDE.md Rule 4).
"""

from __future__ import annotations

import html
import logging
import re
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from weasyprint import HTML

from agent.drafting import fallback_statement
from agent.tools.country_lookup import lookup_country

logger = logging.getLogger("refuge.doc_generator")

TEMPLATES = Path(__file__).resolve().parent / "templates"
_HEAD = (TEMPLATES / "_head.html").read_text(encoding="utf-8")
LOGO_PNG = TEMPLATES / "logo.png"
_PLACEHOLDER = re.compile(r"(\[[^\]]+\])")
# Palette mirrors templates/_head.html so the .docx matches the .pdf exactly.
_AMBER = RGBColor(0xC2, 0x63, 0x29)
_TEAL = RGBColor(0x0A, 0x50, 0x42)
_INK = RGBColor(0x1A, 0x1A, 0x1A)
_MUTED = RGBColor(0x6B, 0x72, 0x80)
_LINE = "E7E2D8"      # light divider (hex, no #)
_TEAL_HEX = "0E6A58"
# Fonts: same families as the PDF. Word uses the named font if installed and
# substitutes otherwise — the .docx still reads as the branded document.
_SERIF = "Fraunces"   # display / headings
_SANS = "Inter"       # body / UI


@dataclass
class GeneratedDoc:
    key: str
    title: str
    meta: str
    pdf_path: Path
    docx_path: Path


# -- shared helpers ---------------------------------------------------------

def _paras(text: str) -> list[str]:
    return [p.strip() for p in text.split("\n") if p.strip()]


def _html_with_placeholders(text: str) -> str:
    """Render free text into HTML, highlighting [placeholders] in amber."""
    out = []
    for para in _paras(text):
        pieces = []
        for part in _PLACEHOLDER.split(para):
            if not part:
                continue
            if _PLACEHOLDER.fullmatch(part):
                pieces.append(f'<span class="fill">{html.escape(part)}</span>')
            else:
                pieces.append(html.escape(part))
        out.append("<p>" + "".join(pieces) + "</p>")
    return "\n".join(out)


# -- DOCX styling layer (mirrors the PDF template in templates/_head.html) -----

def _font(run, name: str, size: float | None = None, color: RGBColor | None = None,
          bold: bool = False, italic: bool = False):
    run.font.name = name
    # Make the font apply to all script ranges, not just Latin.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def _border(paragraph, edge: str, color: str, size: int = 8, space: int = 6):
    """Add a single border on one edge of a paragraph (used for rules/dividers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pbdr.append(el)


def _page_number(paragraph):
    run = paragraph.add_run()
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = f" {text} "
            run._r.append(instr)
    _font(run, _SANS, 8, _MUTED)


def _docx_base(title: str, subtitle: str) -> Document:
    """A new branded document: A4-ish margins, Inter body, the Fugee masthead
    (logo + wordmark + rule), an H1 title, a muted subtitle, and a page footer —
    the same anatomy as the PDF in templates/_head.html."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Mm(22), Mm(18)
    section.left_margin, section.right_margin = Mm(20), Mm(20)

    normal = doc.styles["Normal"]
    normal.font.name = _SANS
    normal.font.size = Pt(11)
    normal.font.color.rgb = _INK
    rfonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), _SANS)

    # Masthead: logo + wordmark on the left, tagline pushed right, teal rule under.
    head = doc.add_paragraph()
    head.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    if LOGO_PNG.exists():
        head.add_run().add_picture(str(LOGO_PNG), height=Inches(0.26))
    _font(head.add_run("  Fugee"), _SERIF, 17, _TEAL, bold=True)
    _font(head.add_run("\tSAFE GUIDANCE FOR PEOPLE ON THE MOVE"), _SANS, 8, _MUTED)
    _border(head, "bottom", _TEAL_HEX, size=18, space=8)
    head.paragraph_format.space_after = Pt(16)

    # Title + subtitle.
    h1 = doc.add_paragraph()
    _font(h1.add_run(title), _SERIF, 22, _TEAL, bold=True)
    h1.paragraph_format.space_after = Pt(2)
    sub = doc.add_paragraph()
    _font(sub.add_run(subtitle), _SANS, 10, _MUTED, italic=True)
    sub.paragraph_format.space_after = Pt(14)

    # Footer: "Prepared with Fugee …  Page N".
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(fp.add_run("Prepared with Fugee · Safe guidance for people on the move    "), _SANS, 8, _MUTED)
    _page_number(fp)
    return doc


def _docx_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    _font(p.add_run(text), _SERIF, 13.5, _TEAL, bold=True)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    _border(p, "bottom", _LINE, size=6, space=3)
    return p


def _docx_body(doc: Document, text: str):
    """Paragraphs with [placeholders] highlighted in amber bold (as in the PDF)."""
    for para in _paras(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        for part in _PLACEHOLDER.split(para):
            if not part:
                continue
            run = p.add_run(part)
            _font(run, _SANS, 11, _INK)
            if _PLACEHOLDER.fullmatch(part):
                _font(run, _SANS, 11, _AMBER, bold=True)


def _docx_bullets(doc: Document, items: list[str]):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _font(p.add_run(str(it)), _SANS, 11, _INK)


def _docx_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    _border(p, "top", _LINE, size=6, space=8)
    _font(p.add_run(text), _SANS, 9.5, _MUTED)


# Branded masthead (the Fugee mark + wordmark) shown atop every PDF.
_LOGO_SVG = (
    '<svg class="logo" viewBox="0 0 32 32" xmlns="http://www.w3.org/2000/svg">'
    '<circle cx="16" cy="16" r="15.5" fill="#0E6A58"/>'
    '<path d="M16 7l7 6.5V25h-4.6v-6.2h-4.8V25H9V13.5L16 7z" fill="#fff"/>'
    '<circle cx="16" cy="13.6" r="1.7" fill="#E07B39"/></svg>'
)
_BRAND = (
    f'<div class="brand">{_LOGO_SVG}<span class="name">Fugee</span>'
    '<span class="tag">Safe guidance for people on the move</span></div>'
)




def _pdf_from_html(body_html: str, path: Path):
    # base_url = the templates dir so @font-face url('fonts/...') resolves locally.
    html_doc = f"<!DOCTYPE html><html><head>{_HEAD}</head><body>{_BRAND}{body_html}</body></html>"
    HTML(string=html_doc, base_url=str(TEMPLATES)).write_pdf(str(path))


# -- per-document content ---------------------------------------------------

def _statement_body_html(text: str) -> str:
    return (
        "<h1>Personal Statement</h1>"
        '<p class="sub">In support of an application for refugee status · Prepared with Fugee</p>'
        f"{_html_with_placeholders(text)}"
        '<p class="note">Fields in amber are placeholders for you to complete. You can edit any part '
        "of this statement in the Word (.docx) version.</p>"
    )


def _rec_for(session):
    if session.selected_country:
        rec = lookup_country(session.selected_country)
        return None if rec.get("error") else rec
    return None


def _action_plan_blocks(session, rec):
    country = session.selected_country or "your chosen country"
    months = (rec or {}).get("processingTimeMonths")
    office = (rec or {}).get("unhcrOffice")
    steps = (rec or {}).get("steps") or [
        "Register with UNHCR or the asylum authority",
        "File your asylum claim with your personal statement",
        "Attend your refugee status interview (RSD)",
        "Receive the decision (appeal if refused)",
        "Access integration support",
    ]
    intro = (f"This plan is for seeking protection in {country}. "
             f"Typical processing: {('about ' + str(months) + ' months') if months else 'varies'}. "
             f"UNHCR: {('Yes — ' + office) if office else 'check locally'}.")
    return country, intro, steps


def _orgs(rec):
    return [(o.get("name", ""), o.get("url", "")) for o in (rec or {}).get("legalAidOrgs", []) if isinstance(o, dict)]


# -- generation -------------------------------------------------------------

def generate(session, statement: str | None = None, out_dir: Path | None = None) -> list[GeneratedDoc]:
    """Build the 4 documents as PDF + DOCX. ``statement`` is the LLM-drafted
    personal statement; if omitted, a deterministic fallback is used."""
    out_dir = Path(out_dir) if out_dir else Path(tempfile.mkdtemp(prefix="refuge_docs_"))
    out_dir.mkdir(parents=True, exist_ok=True)
    stmt = statement or fallback_statement(session)
    rec = _rec_for(session)
    docs: list[GeneratedDoc] = []

    # 1. Personal statement (LLM-drafted) — PDF + DOCX
    ps_pdf, ps_docx = out_dir / "personal_statement.pdf", out_dir / "personal_statement.docx"
    _pdf_from_html(_statement_body_html(stmt), ps_pdf)
    d = _docx_base("Personal Statement", "In support of an application for refugee status · Prepared with Fugee")
    _docx_body(d, stmt)
    _docx_note(d, "Fields in amber are placeholders for you to complete. You can edit "
                  "any part of this statement in this Word (.docx) version.")
    d.save(str(ps_docx))
    docs.append(GeneratedDoc("personal_statement", "Personal statement (editable)",
                             "Word + PDF · drafted for you, with placeholders to complete", ps_pdf, ps_docx))

    # 2. Action plan
    country, intro, steps = _action_plan_blocks(session, rec)
    ap_html = (f"<h1>Action Plan — {html.escape(country)}</h1>"
               '<p class="sub">Step-by-step roadmap · Prepared with Fugee</p>'
               f"<p>{html.escape(intro)}</p>" +
               "".join(f'<p class="step"><b>Step {i}.</b> {html.escape(str(s))}</p>'
                       for i, s in enumerate(steps, 1)))
    ap_pdf, ap_docx = out_dir / "action_plan.pdf", out_dir / "action_plan.docx"
    _pdf_from_html(ap_html, ap_pdf)
    d = _docx_base(f"Action Plan — {country}", "Step-by-step roadmap · Prepared with Fugee")
    _docx_body(d, intro)
    for i, s in enumerate(steps, 1):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _font(p.add_run(f"Step {i}.  "), _SANS, 11, _TEAL, bold=True)
        _font(p.add_run(str(s)), _SANS, 11, _INK)
    d.save(str(ap_docx))
    docs.append(GeneratedDoc("action_plan", "Action plan", "Word + PDF · roadmap with contacts", ap_pdf, ap_docx))

    # 3. Emergency contacts
    office = (rec or {}).get("unhcrOffice") or "Nearest UNHCR office"
    orgs = _orgs(rec)
    ec_html = ("<h1>Emergency Contacts</h1>"
               '<p class="sub">UNHCR offices &amp; legal aid · Prepared with Fugee</p>'
               f"<h2>{html.escape(country)}</h2><p>UNHCR office: {html.escape(office)}</p>"
               "<h2>Legal aid &amp; support</h2>" +
               ("".join(f'<div class="org"><b>{html.escape(n)}</b><span>{html.escape(u)}</span></div>' for n, u in orgs)
                or "<p>Ask the UNHCR office for registered legal-aid partners.</p>"))
    ec_pdf, ec_docx = out_dir / "emergency_contacts.pdf", out_dir / "emergency_contacts.docx"
    _pdf_from_html(ec_html, ec_pdf)
    d = _docx_base("Emergency Contacts", "UNHCR offices & legal aid · Prepared with Fugee")
    _docx_h2(d, country)
    _docx_body(d, f"UNHCR office: {office}")
    _docx_h2(d, "Legal aid & support")
    if orgs:
        for n, u in orgs:
            p = d.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _font(p.add_run(n + "\n"), _SANS, 11, _TEAL, bold=True)
            _font(p.add_run(u), _SANS, 10, _MUTED)
    else:
        _docx_body(d, "Ask the UNHCR office for registered legal-aid partners.")
    d.save(str(ec_docx))
    docs.append(GeneratedDoc("emergency_contacts", "Emergency contacts", "Word + PDF · UNHCR & legal aid", ec_pdf, ec_docx))

    # 4. Rights summary card
    grounds = ", ".join(session.assessment.convention_grounds) if session.assessment.convention_grounds else "to be confirmed"
    rc_html = ("<h1>Your Rights — Summary Card</h1>"
               '<p class="sub">Key protections · Prepared with Fugee</p>'
               "<h2>Non-refoulement</h2><p>You cannot be forced back to a country where your life or freedom "
               "would be at serious risk (1951 Refugee Convention).</p>"
               "<h2>While your claim is decided</h2><ul>"
               "<li>The right to seek asylum and a fair hearing.</li>"
               "<li>The right to an interpreter and free legal assistance.</li>"
               "<li>The right not to be detained arbitrarily.</li>"
               "<li>The right to shelter, food, and emergency healthcare.</li></ul>"
               f"<p>Origin: {html.escape(session.interview.origin_country or '[origin]')} · "
               f"Seeking protection in: {html.escape(country)} · Grounds: {html.escape(grounds)}</p>")
    rc_pdf, rc_docx = out_dir / "rights_summary_card.pdf", out_dir / "rights_summary_card.docx"
    _pdf_from_html(rc_html, rc_pdf)
    d = _docx_base("Your Rights — Summary Card", "Key protections · Prepared with Fugee")
    _docx_h2(d, "Non-refoulement")
    _docx_body(d, "You cannot be forced back to a country where your life or freedom "
                  "would be at serious risk (1951 Refugee Convention).")
    _docx_h2(d, "While your claim is decided")
    _docx_bullets(d, ["The right to seek asylum and a fair hearing.",
                      "The right to an interpreter and free legal assistance.",
                      "The right not to be detained arbitrarily.",
                      "The right to shelter, food, and emergency healthcare."])
    _docx_body(d, f"Origin: {session.interview.origin_country or '[origin]'} · "
                  f"Seeking protection in: {country} · Grounds: {grounds}")
    d.save(str(rc_docx))
    docs.append(GeneratedDoc("rights_summary_card", "Your rights — summary card", "Word + PDF · key protections", rc_pdf, rc_docx))

    logger.info("Generated %d documents (PDF + DOCX) in %s", len(docs), out_dir)
    return docs


def all_files(docs: list[GeneratedDoc]) -> list[str]:
    files = []
    for d in docs:
        files.append(str(d.docx_path))
        files.append(str(d.pdf_path))
    return files


def zip_package(docs: list[GeneratedDoc], out_dir: Path | None = None) -> Path:
    base = Path(out_dir) if out_dir else docs[0].pdf_path.parent
    zip_path = base / "refuge_documents.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in all_files(docs):
            zf.write(f, arcname=Path(f).name)
    return zip_path


def preview_statement_html(session, statement: str | None = None) -> str:
    """Head-less personal-statement snippet for the on-screen preview."""
    stmt = statement or fallback_statement(session)
    return ('<article class="doc"><h4>Personal Statement</h4>'
            '<p class="doc__sub">In support of an application for refugee status · Prepared with Fugee</p>'
            f"{_html_with_placeholders(stmt)}</article>")


__all__ = ["generate", "preview_statement_html", "zip_package", "all_files", "GeneratedDoc"]
